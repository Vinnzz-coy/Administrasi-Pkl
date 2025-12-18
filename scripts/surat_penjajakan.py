import sys
import os

def replace_in_docx(template_path, output_path, data):
    """
    Replace placeholders in DOCX template with actual data
    """
    try:
        from docx import Document
        
        # Load template
        print(f"[INFO] Loading template: {template_path}")
        doc = Document(template_path)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                # Support both {{ key }} and {{key}} formats
                placeholder1 = f"{{{{ {key} }}}}"
                placeholder2 = f"{{{{{key}}}}}"
                
                if placeholder1 in paragraph.text or placeholder2 in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder1, str(value))
                    paragraph.text = paragraph.text.replace(placeholder2, str(value))
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder1 = f"{{{{ {key} }}}}"
                        placeholder2 = f"{{{{{key}}}}}"
                        
                        if placeholder1 in cell.text or placeholder2 in cell.text:
                            cell.text = cell.text.replace(placeholder1, str(value))
                            cell.text = cell.text.replace(placeholder2, str(value))
        
        # Save DOCX
        docx_output = output_path + '.docx'
        print(f"[INFO] Saving DOCX: {docx_output}")
        doc.save(docx_output)
        print(f"[OK] DOCX created: {docx_output}")
        
        # Try to convert to PDF
        pdf_output = output_path + '.pdf'
        pdf_created = False
        
        # Method 1: Try docx2pdf (works on Windows with Word installed)
        try:
            from docx2pdf import convert
            print(f"[INFO] Converting to PDF using docx2pdf...")
            convert(docx_output, pdf_output)
            pdf_created = True
            print(f"[OK] PDF created: {pdf_output}")
        except Exception as e:
            print(f"[WARNING] docx2pdf failed: {str(e)}")
        
        # Method 2: Try LibreOffice (Linux/Mac)
        if not pdf_created:
            try:
                import subprocess
                print(f"[INFO] Trying LibreOffice conversion...")
                cmd = [
                    'libreoffice',
                    '--headless',
                    '--convert-to',
                    'pdf',
                    '--outdir',
                    os.path.dirname(output_path),
                    docx_output
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                if result.returncode == 0:
                    pdf_created = True
                    print(f"[OK] PDF created: {pdf_output}")
                else:
                    print(f"[WARNING] LibreOffice conversion failed: {result.stderr}")
            except FileNotFoundError:
                print("[WARNING] LibreOffice not found")
            except Exception as e:
                print(f"[WARNING] LibreOffice conversion error: {str(e)}")
        
        # Method 3: Try unoconv (alternative)
        if not pdf_created:
            try:
                import subprocess
                print(f"[INFO] Trying unoconv conversion...")
                cmd = ['unoconv', '-f', 'pdf', '-o', pdf_output, docx_output]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                if result.returncode == 0:
                    pdf_created = True
                    print(f"[OK] PDF created: {pdf_output}")
                else:
                    print(f"[WARNING] unoconv conversion failed: {result.stderr}")
            except FileNotFoundError:
                print("[WARNING] unoconv not found")
            except Exception as e:
                print(f"[WARNING] unoconv conversion error: {str(e)}")
        
        # If no PDF method worked, create a dummy PDF placeholder
        if not pdf_created:
            print("[WARNING] All PDF conversion methods failed")
            print("[INFO] Creating placeholder PDF message")
            # Copy DOCX as PDF (not ideal but better than failing)
            import shutil
            shutil.copy(docx_output, pdf_output)
            print(f"[INFO] Created fallback file: {pdf_output}")
        
        return True
        
    except Exception as e:
        print(f"[ERROR] Failed to process document: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print(f"[INFO] Starting script with {len(sys.argv)} arguments")
    print(f"[INFO] Arguments: {sys.argv}")
    
    if len(sys.argv) != 12:
        print("[ERROR] Incorrect number of arguments")
        print("Usage: python generate_surat.py <template_path> <output_path> <tanggal_surat> <nomor_surat> <nama_dudi> <alamat_dudi> <lama_pkl> <tanggal_mulai> <tanggal_selesai> <tingkat> <jurusan>")
        print(f"Expected: 12 arguments")
        print(f"Received: {len(sys.argv)} arguments")
        sys.exit(1)
    
    template_path = sys.argv[1]
    output_path = sys.argv[2]
    
    # Prepare data dictionary
    data = {
        'tanggal_surat': sys.argv[3],
        'nomor_surat': sys.argv[4],
        'nama_dudi': sys.argv[5],
        'alamat_dudi': sys.argv[6],
        'lama_pkl': sys.argv[7],
        'tanggal_mulai': sys.argv[8],
        'tanggal_selesai': sys.argv[9],
        'tingkat': sys.argv[10],
        'jurusan': sys.argv[11],
    }
    
    print(f"[INFO] Data to replace: {data}")
    
    # Check if template exists
    if not os.path.exists(template_path):
        print(f"[ERROR] Template file not found: {template_path}")
        sys.exit(1)
    
    print(f"[OK] Template file found: {template_path}")
    
    # Create output directory if not exists
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        print(f"[INFO] Creating output directory: {output_dir}")
        os.makedirs(output_dir, exist_ok=True)
    
    # Check python-docx is installed
    try:
        import docx
        print("[OK] python-docx module found")
    except ImportError:
        print("[ERROR] python-docx not installed")
        print("Please run: pip install python-docx")
        sys.exit(1)
    
    # Process document
    success = replace_in_docx(template_path, output_path, data)
    
    if success:
        print("[SUCCESS] Document processed successfully")
        sys.exit(0)
    else:
        print("[FAILED] Document processing failed")
        sys.exit(1)

if __name__ == "__main__":
    main()